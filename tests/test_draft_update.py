from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document
from PIL import Image

from wxcli.draft_update import DraftUpdatePlanner
from wxcli.errors import ValidationError
from wxcli.official_draft import DraftSnapshot


class FakeWriter:
    def __init__(self) -> None:
        self.snapshot_value = DraftSnapshot(
            media_id="draft-id",
            news_items=[{
                "title": "旧标题",
                "content": '<p>旧正文</p><img src="https://old.example/image.jpg">',
                "author": "旧作者",
            }],
            fingerprint="remote-fingerprint",
        )
        self.updated: tuple[object, ...] | None = None

    def snapshot(self, media_id: str) -> DraftSnapshot:
        assert media_id == "draft-id"
        return self.snapshot_value

    def update(self, *args: object) -> object:
        self.updated = args
        return {"updated": True}


def _inputs(tmp_path: Path) -> tuple[Path, Path]:
    image_path = tmp_path / "body.png"
    cover = tmp_path / "cover.png"
    Image.new("RGB", (640, 360), "red").save(image_path)
    Image.new("RGB", (1200, 510), "blue").save(cover)
    source = tmp_path / "article.docx"
    document = Document()
    document.add_heading("新标题", level=0)
    document.add_paragraph("新正文")
    document.add_picture(str(image_path))
    document.save(source)
    return source, cover


def test_backup_never_overwrites_and_plan_contains_diff_and_snapshot(tmp_path: Path) -> None:
    writer = FakeWriter()
    planner = DraftUpdatePlanner(writer)  # type: ignore[arg-type]
    backup = tmp_path / "backup.json"

    result = planner.backup("draft-id", backup)
    assert result.fingerprint == "remote-fingerprint"
    assert json.loads(backup.read_text(encoding="utf-8"))["news_items"][0]["title"] == "旧标题"
    with pytest.raises(ValidationError, match="already exists"):
        planner.backup("draft-id", backup)

    source, cover = _inputs(tmp_path)
    plan = planner.plan("draft-id", 0, source, cover, tmp_path / "plan")
    assert plan.difference.changed is True
    assert set(plan.difference.changes) >= {"title", "body_text", "body_images"}
    assert "author" not in plan.difference.changes
    assert (tmp_path / "plan" / "backup.json").is_file()
    assert (tmp_path / "plan" / "prepared" / "package.json").is_file()


def test_update_requires_confirmation_and_uses_frozen_fingerprint(tmp_path: Path) -> None:
    writer = FakeWriter()
    planner = DraftUpdatePlanner(writer)  # type: ignore[arg-type]
    source, cover = _inputs(tmp_path)
    plan_dir = tmp_path / "plan"
    planner.plan("draft-id", 0, source, cover, plan_dir)

    with pytest.raises(ValidationError, match="requires --confirm"):
        planner.apply(plan_dir, confirmed=False)
    result = planner.apply(plan_dir, confirmed=True)

    assert result == {"updated": True}
    assert writer.updated is not None
    assert writer.updated[0:2] == ("draft-id", 0)
    assert writer.updated[3] == "remote-fingerprint"
