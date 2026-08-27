from __future__ import annotations

import hashlib
import json
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from PIL import Image
from typer.testing import CliRunner

from wxcli.cli import app
from wxcli.draft_import import WordDraftImporter
from wxcli.errors import ValidationError


def _noise_png(path: Path, size: tuple[int, int]) -> None:
    image = Image.effect_noise(size, 80).convert("RGB")
    image.save(path, format="PNG")


def _word_article(path: Path, image_path: Path, *, with_table: bool = False) -> None:
    document = Document()
    document.add_heading("示例标题", level=0)
    document.sections[0].header.paragraphs[0].text = "不要导入页眉"
    document.sections[0].footer.paragraphs[0].text = "不要导入页脚"
    document.add_paragraph("第一段正文")
    document.add_picture(str(image_path))
    document.add_heading("金色小标题", level=1)
    document.add_paragraph("第二段正文")
    if with_table:
        document.add_table(rows=1, cols=1).cell(0, 0).text = "表格"
    document.save(path)


def _add_hyperlink(paragraph: object, label: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(  # type: ignore[attr-defined]
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)  # type: ignore[attr-defined]


def test_word_draft_import_preserves_order_and_prepares_wechat_images(tmp_path: Path) -> None:
    source_image = tmp_path / "body.png"
    cover = tmp_path / "cover.png"
    source = tmp_path / "article.docx"
    output = tmp_path / "preview"
    _noise_png(source_image, (1920, 1080))
    _noise_png(cover, (1200, 510))
    _word_article(source, source_image)
    source_hash = hashlib.sha256(source.read_bytes()).hexdigest()
    cover_hash = hashlib.sha256(cover.read_bytes()).hexdigest()

    prepared = WordDraftImporter().prepare(source, cover, output)

    assert prepared.title == "示例标题"
    assert len(prepared.images) == 1
    assert prepared.images[0].prepared_bytes < 1_000_000
    assert prepared.cover.prepared_bytes < 64_000
    assert prepared.content_template.index("第一段正文") < prepared.content_template.index(
        "wxcli-image-001"
    )
    assert prepared.content_template.index("wxcli-image-001") < prepared.content_template.index(
        "金色小标题"
    )
    assert "#9a6a22" in prepared.content_template
    assert "不要导入页眉" not in prepared.content_template
    assert "不要导入页脚" not in prepared.content_template
    assert Path(prepared.preview.preview_html).is_file()
    assert Path(prepared.preview.manifest).is_file()
    manifest = json.loads(Path(prepared.preview.manifest).read_text(encoding="utf-8"))
    assert manifest["ready_for_upload"] is True
    assert manifest["content_image_count"] == 1
    assert hashlib.sha256(source.read_bytes()).hexdigest() == source_hash
    assert hashlib.sha256(cover.read_bytes()).hexdigest() == cover_hash


def test_uploaded_urls_replace_every_placeholder(tmp_path: Path) -> None:
    source_image = tmp_path / "body.png"
    cover = tmp_path / "cover.png"
    source = tmp_path / "article.docx"
    _noise_png(source_image, (640, 360))
    _noise_png(cover, (1200, 510))
    _word_article(source, source_image)
    prepared = WordDraftImporter().prepare(source, cover, tmp_path / "preview")

    content = prepared.content_with_urls(["https://mmbiz.qpic.cn/example/0"])

    assert "wxcli-image-001" not in content
    assert 'src="https://mmbiz.qpic.cn/example/0"' in content


def test_nonempty_output_is_rejected_and_simple_tables_are_mapped(tmp_path: Path) -> None:
    source_image = tmp_path / "body.png"
    cover = tmp_path / "cover.png"
    source = tmp_path / "article.docx"
    _noise_png(source_image, (640, 360))
    _noise_png(cover, (1200, 510))
    _word_article(source, source_image, with_table=True)
    output = tmp_path / "preview"
    output.mkdir()
    (output / "existing.txt").write_text("keep", encoding="utf-8")

    with pytest.raises(ValidationError, match="must be empty"):
        WordDraftImporter().prepare(source, cover, output)
    assert (output / "existing.txt").read_text(encoding="utf-8") == "keep"

    prepared = WordDraftImporter().prepare(source, cover, tmp_path / "empty")
    assert "<table" in prepared.content_template
    assert "表格" in prepared.content_template


def test_prepared_package_detects_image_tampering(tmp_path: Path) -> None:
    source_image = tmp_path / "body.png"
    cover = tmp_path / "cover.png"
    source = tmp_path / "article.docx"
    output = tmp_path / "preview"
    _noise_png(source_image, (640, 360))
    _noise_png(cover, (1200, 510))
    _word_article(source, source_image)
    prepared = WordDraftImporter().prepare(source, cover, output)

    loaded = type(prepared).load(output)
    assert loaded.package_sha256 == prepared.package_sha256
    loaded.images[0].path.write_bytes(b"changed")
    with pytest.raises(ValidationError, match="changed"):
        type(prepared).load(output)


def test_richer_word_layout_mapping(tmp_path: Path) -> None:
    source_image = tmp_path / "body.png"
    cover = tmp_path / "cover.png"
    source = tmp_path / "rich.docx"
    _noise_png(source_image, (640, 360))
    _noise_png(cover, (1200, 510))
    document = Document()
    document.add_heading("完整映射", level=0)
    document.add_heading("二级", level=2)
    document.add_heading("三级", level=3)
    linked = document.add_paragraph("链接：")
    _add_hyperlink(linked, "微信文档", "https://developers.weixin.qq.com/")
    document.add_paragraph("项目", style="List Bullet")
    document.add_paragraph("步骤", style="List Number")
    document.add_paragraph("引用内容", style="Quote")
    document.add_paragraph("图片说明", style="Caption")
    centered = document.add_paragraph("居中内容")
    centered.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mixed = document.add_paragraph("图前")
    mixed.add_run().add_picture(str(source_image))
    mixed.add_run("图后")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "左格"
    table.cell(0, 1).text = "右格"
    document.save(source)

    prepared = WordDraftImporter().prepare(source, cover, tmp_path / "rich-preview")
    content = prepared.content_template

    assert "<h3" in content and "二级" in content
    assert "<h4" in content and "三级" in content
    assert '<a href="https://developers.weixin.qq.com/">微信文档</a>' in content
    assert "<ul><li>项目</li></ul>" in content
    assert "<ol><li>步骤</li></ol>" in content
    assert "<blockquote" in content and "引用内容" in content
    assert "font-size:14px" in content and "图片说明" in content
    assert "text-align:center" in content and "居中内容" in content
    assert content.index("图前") < content.index("wxcli-image-001") < content.index("图后")
    assert "<table" in content and "左格" in content and "右格" in content


def test_cli_preview_is_local_and_emits_one_json_document(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source_image = tmp_path / "body.png"
    cover = tmp_path / "cover.png"
    source = tmp_path / "article.docx"
    output = tmp_path / "preview"
    _noise_png(source_image, (640, 360))
    _noise_png(cover, (1200, 510))
    _word_article(source, source_image)

    def reject_network(*args: object, **kwargs: object) -> None:
        pytest.fail("local preview must not create an HTTP client")

    monkeypatch.setattr("wxcli.cli.httpx.Client", reject_network)
    result = CliRunner().invoke(
        app,
        [
            "--json",
            "account",
            "draft",
            "import-word",
            str(source),
            "--cover",
            str(cover),
            "--output",
            str(output),
        ],
    )

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    assert payload["ok"] is True
    assert payload["data"]["ready_for_upload"] is True
    assert result.stderr == ""
