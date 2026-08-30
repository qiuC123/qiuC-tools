"""Prepare a Word document and cover image for safe Official Account draft creation."""

from __future__ import annotations

import html
import hashlib
import json
import shutil
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile

from docx import Document
from docx.document import Document as DocumentObject
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from PIL import Image, ImageOps, UnidentifiedImageError
from pydantic import BaseModel, ConfigDict, Field

from wxcli.errors import ErrorCode, ValidationError, WxcliError

_BODY_IMAGE_LIMIT = 1_000_000
_COVER_IMAGE_LIMIT = 64_000
_CONTENT_CHARACTER_LIMIT = 20_000
_CONTENT_BYTE_LIMIT = 1_000_000


class DraftImportPreview(BaseModel):
    """Non-secret facts about a locally prepared draft."""

    model_config = ConfigDict(extra="forbid")

    title: str
    source_docx: str
    cover_source: str
    preview_html: str
    manifest: str
    package_manifest: str = ""
    package_sha256: str = ""
    content_image_count: int = Field(ge=0)
    original_image_bytes: int = Field(ge=0)
    prepared_image_bytes: int = Field(ge=0)
    content_characters: int = Field(ge=0)
    ready_for_upload: bool = True


@dataclass(frozen=True, slots=True)
class PreparedImage:
    """A validated JPEG ready for one controlled WeChat upload."""

    path: Path
    original_bytes: int
    prepared_bytes: int
    width: int
    height: int
    sha256: str = ""


@dataclass(frozen=True, slots=True)
class PreparedDraft:
    """A local draft package containing no credentials or remote state."""

    title: str
    author: str | None
    digest: str | None
    content_template: str
    images: tuple[PreparedImage, ...]
    cover: PreparedImage
    preview: DraftImportPreview

    @property
    def package_sha256(self) -> str:
        """Stable identity for the exact text and prepared assets."""
        value = {
            "title": self.title,
            "author": self.author,
            "digest": self.digest,
            "content_template": self.content_template,
            "images": [image.sha256 for image in self.images],
            "cover": self.cover.sha256,
        }
        encoded = json.dumps(value, ensure_ascii=False, sort_keys=True).encode("utf-8")
        return hashlib.sha256(encoded).hexdigest()

    def content_with_urls(self, urls: list[str]) -> str:
        """Replace every local image placeholder with one uploaded WeChat URL."""
        if len(urls) != len(self.images):
            raise ValidationError("The uploaded image count does not match the prepared draft.")
        content = self.content_template
        for index, url in enumerate(urls, start=1):
            content = content.replace(_placeholder(index), html.escape(url, quote=True))
        _validate_content(content)
        return content

    @classmethod
    def load(cls, output_dir: Path) -> PreparedDraft:
        """Load and integrity-check a package that was previously previewed."""
        root = output_dir.expanduser().resolve()
        package_path = root / "package.json"
        try:
            value = json.loads(package_path.read_text(encoding="utf-8"))
        except (OSError, ValueError) as error:
            raise ValidationError("The prepared draft package could not be read.") from error
        if not isinstance(value, dict) or value.get("schema_version") != 1:
            raise ValidationError("The prepared draft package format is invalid.")
        try:
            image_values = value["images"]
            if not isinstance(image_values, list):
                raise TypeError
            images = tuple(_load_prepared_image(root, item) for item in image_values)
            cover = _load_prepared_image(root, value["cover"])
            preview = DraftImportPreview.model_validate(value["preview"])
            prepared = cls(
                title=str(value["title"]),
                author=_string_or_none(value.get("author")),
                digest=_string_or_none(value.get("digest")),
                content_template=str(value["content_template"]),
                images=images,
                cover=cover,
                preview=preview,
            )
        except (KeyError, TypeError, ValueError) as error:
            raise ValidationError("The prepared draft package is invalid or has changed.") from error
        _validate_content(prepared.content_template)
        if prepared.package_sha256 != value.get("package_sha256"):
            raise ValidationError("The prepared draft package has changed since preview.")
        return prepared


class WordDraftImporter:
    """Convert a simple Word article into WeChat-safe HTML and compressed images."""

    def prepare(
        self,
        source: Path,
        cover: Path,
        output_dir: Path,
        *,
        author: str | None = None,
        digest: str | None = None,
    ) -> PreparedDraft:
        source = _validate_file(source, ".docx", "Word document")
        cover = _validate_image_file(cover, "cover image")
        author = _optional_limited_text(author, 16, "author")
        digest = _optional_limited_text(digest, 120, "digest")
        _validate_empty_output_directory(output_dir)

        try:
            document = Document(str(source))
        except (OSError, ValueError, BadZipFile, PackageNotFoundError) as error:
            raise WxcliError(
                ErrorCode.PARSING_ERROR,
                "The Word document could not be opened.",
            ) from error
        title_paragraph = self._title_paragraph(document.paragraphs)
        title = title_paragraph.text.strip()
        if len(title) > 32:
            raise ValidationError("The draft title must not exceed 32 characters.")

        blocks: list[str] = []
        preview_blocks: list[str] = []
        image_blobs: list[bytes] = []
        for item in _document_blocks(document):
            if isinstance(item, Table):
                table_html = self._table_block(item)
                blocks.append(table_html)
                preview_blocks.append(table_html)
                continue
            paragraph = item
            if paragraph._p is title_paragraph._p:
                continue
            text = paragraph.text.strip()
            paragraph_parts = self._paragraph_parts(paragraph)
            relationship_ids = self._image_relationship_ids(paragraph)
            if relationship_ids:
                fragments: list[str] = []
                preview_fragments: list[str] = []
                for kind, value in paragraph_parts:
                    if kind == "html":
                        fragments.append(value)
                        preview_fragments.append(value)
                        continue
                    relationship_id = value
                    part = document.part.related_parts.get(relationship_id)
                    blob = getattr(part, "blob", None)
                    if not isinstance(blob, bytes):
                        raise WxcliError(
                            ErrorCode.PARSING_ERROR,
                            "The Word document contains an unreadable image.",
                        )
                    image_blobs.append(blob)
                    index = len(image_blobs)
                    image_style = "max-width:100%;height:auto"
                    fragments.append(f'<img src="{_placeholder(index)}" style="{image_style}">')
                    relative = f"images/body-{index:03d}.jpg"
                    preview_fragments.append(
                        f'<img src="{html.escape(relative, quote=True)}" '
                        f'style="{image_style}">'
                    )
                blocks.append(self._wrap_paragraph(paragraph, "".join(fragments)))
                preview_blocks.append(self._wrap_paragraph(paragraph, "".join(preview_fragments)))
                continue
            if not text and not paragraph_parts:
                continue
            block = self._text_block(paragraph)
            blocks.append(block)
            preview_blocks.append(block)

        if not blocks:
            raise ValidationError("The Word document does not contain draft body content.")

        content_template = self._content_document(blocks)
        _validate_content(content_template)
        preview_content = self._content_document(preview_blocks)
        output_existed = output_dir.exists()
        _prepare_empty_directory(output_dir)
        image_dir = output_dir / "images"
        image_dir.mkdir()
        try:
            prepared_images = [
                _prepare_jpeg(
                    blob,
                    image_dir / f"body-{index:03d}.jpg",
                    _BODY_IMAGE_LIMIT,
                )
                for index, blob in enumerate(image_blobs, start=1)
            ]
            cover_target = image_dir / "cover.jpg"
            cover_prepared = _prepare_jpeg(
                cover.read_bytes(),
                cover_target,
                _COVER_IMAGE_LIMIT,
            )
            preview_path = output_dir / "preview.html"
            preview_path.write_text(
                self._preview_document(title, preview_content),
                encoding="utf-8",
            )

            original_image_bytes = sum(image.original_bytes for image in prepared_images)
            prepared_image_bytes = sum(image.prepared_bytes for image in prepared_images)
            manifest_path = output_dir / "manifest.json"
            package_path = output_dir / "package.json"
            provisional_preview = DraftImportPreview(
                title=title,
                source_docx=str(source),
                cover_source=str(cover),
                preview_html=str(preview_path),
                manifest=str(manifest_path),
                package_manifest=str(package_path),
                package_sha256="0" * 64,
                content_image_count=len(prepared_images),
                original_image_bytes=original_image_bytes,
                prepared_image_bytes=prepared_image_bytes,
                content_characters=len(content_template),
            )
            provisional = PreparedDraft(
                title=title,
                author=author,
                digest=digest,
                content_template=content_template,
                images=tuple(prepared_images),
                cover=cover_prepared,
                preview=provisional_preview,
            )
            preview = DraftImportPreview(
                title=title,
                source_docx=str(source),
                cover_source=str(cover),
                preview_html=str(preview_path),
                manifest=str(manifest_path),
                package_manifest=str(package_path),
                package_sha256=provisional.package_sha256,
                content_image_count=len(prepared_images),
                original_image_bytes=original_image_bytes,
                prepared_image_bytes=prepared_image_bytes,
                content_characters=len(content_template),
            )
            manifest_path.write_text(
                json.dumps(preview.model_dump(mode="json"), ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            prepared = PreparedDraft(
                title=title,
                author=author,
                digest=digest,
                content_template=content_template,
                images=tuple(prepared_images),
                cover=cover_prepared,
                preview=preview,
            )
            package_path.write_text(
                json.dumps(_package_value(prepared, output_dir), ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            return prepared
        except Exception:
            _clean_failed_output(output_dir, remove_root=not output_existed)
            raise

    @staticmethod
    def _title_paragraph(paragraphs: list[Paragraph]) -> Paragraph:
        for paragraph in paragraphs:
            if (
                paragraph.text.strip()
                and paragraph.style is not None
                and paragraph.style.name == "Title"
            ):
                return paragraph
        for paragraph in paragraphs:
            if paragraph.text.strip():
                return paragraph
        raise ValidationError("The Word document does not contain a title.")

    @staticmethod
    def _image_relationship_ids(paragraph: Paragraph) -> list[str]:
        if paragraph._p.xpath(".//wp:anchor"):
            raise ValidationError("Floating Word images are not supported by draft import.")
        values: list[str] = []
        for blip in paragraph._p.xpath(".//a:blip"):
            relationship_id = blip.get(qn("r:embed"))
            if relationship_id:
                values.append(str(relationship_id))
        return values

    @staticmethod
    def _text_block(paragraph: Paragraph) -> str:
        parts = WordDraftImporter._paragraph_parts(paragraph)
        text = "".join(value for kind, value in parts if kind == "html") or _formatted_text(paragraph)
        style_name = paragraph.style.name if paragraph.style is not None else ""
        heading = {"Heading 1": ("h2", 20), "Heading 2": ("h3", 18), "Heading 3": ("h4", 17)}.get(style_name)
        if heading is not None:
            tag, size = heading
            return (
                f'<{tag} style="color:#9a6a22;font-size:{size}px">'
                f"{text}</{tag}>"
            )
        return WordDraftImporter._wrap_paragraph(paragraph, text)

    @staticmethod
    def _wrap_paragraph(paragraph: Paragraph, text: str) -> str:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if "Quote" in style_name:
            return f'<blockquote style="margin:16px 8px;padding-left:12px;border-left:3px solid #c49a51;color:#666">{text}</blockquote>'
        if style_name == "Caption":
            return f'<p style="text-align:center;font-size:14px;color:#888">{text}</p>'
        list_tag = _list_tag(paragraph, style_name)
        if list_tag is not None:
            return f"<{list_tag}><li>{text}</li></{list_tag}>"
        alignment = getattr(paragraph.alignment, "name", None)
        css = (
            {"CENTER": "center", "RIGHT": "right", "JUSTIFY": "justify"}.get(alignment)
            if isinstance(alignment, str)
            else None
        )
        attribute = f' style="text-align:{css}"' if css else ""
        return f"<p{attribute}>{text}</p>"

    @staticmethod
    def _paragraph_parts(paragraph: Paragraph) -> list[tuple[str, str]]:
        parts: list[tuple[str, str]] = []
        for child in paragraph._p.iterchildren():
            if child.tag == qn("w:r"):
                run = Run(child, paragraph)
                if run.text:
                    parts.append(("html", _formatted_run(run)))
                for blip in child.xpath(".//a:blip"):
                    relationship_id = blip.get(qn("r:embed"))
                    if relationship_id:
                        parts.append(("image", str(relationship_id)))
            elif child.tag == qn("w:hyperlink"):
                relationship_id = child.get(qn("r:id"))
                label = "".join(str(node.text or "") for node in child.xpath(".//w:t"))
                if not label:
                    continue
                url = None
                if relationship_id:
                    relationship = paragraph.part.rels.get(str(relationship_id))
                    url = getattr(relationship, "target_ref", None)
                safe_url = _safe_link(url)
                escaped = html.escape(label)
                parts.append(("html", f'<a href="{html.escape(safe_url, quote=True)}">{escaped}</a>' if safe_url else escaped))
        return parts

    @staticmethod
    def _table_block(table: Table) -> str:
        if table._tbl.xpath(".//w:gridSpan") or table._tbl.xpath(".//w:vMerge"):
            raise ValidationError("Merged Word table cells are not supported by draft import.")
        rows: list[str] = []
        for row in table.rows:
            cells: list[str] = []
            for cell in row.cells:
                if cell.tables:
                    raise ValidationError("Nested Word tables are not supported by draft import.")
                if any(WordDraftImporter._image_relationship_ids(p) for p in cell.paragraphs):
                    raise ValidationError("Images inside Word table cells are not supported by draft import.")
                value = "<br>".join(_formatted_text(p) for p in cell.paragraphs if p.text.strip())
                cells.append(f'<td style="border:1px solid #ddd;padding:6px">{value}</td>')
            rows.append(f"<tr>{''.join(cells)}</tr>")
        return '<table style="width:100%;border-collapse:collapse">' + "".join(rows) + "</table>"

    @staticmethod
    def _content_document(blocks: list[str]) -> str:
        compacted: list[str] = []
        normal_paragraphs: list[str] = []

        def flush_normal_paragraphs() -> None:
            if normal_paragraphs:
                compacted.append(f"<p>{'<br><br>'.join(normal_paragraphs)}</p>")
                normal_paragraphs.clear()

        for block in blocks:
            if block.startswith("<p>") and block.endswith("</p>"):
                normal_paragraphs.append(block[3:-4])
                continue
            flush_normal_paragraphs()
            compacted.append(block)
        flush_normal_paragraphs()
        return (
            '<section style="font-size:17px;line-height:1.9;color:#222;'
            'text-align:justify;word-break:break-word">'
            f"{''.join(compacted)}</section>"
        )

    @staticmethod
    def _preview_document(title: str, content: str) -> str:
        return (
            "<!doctype html><html lang=\"zh-CN\"><head><meta charset=\"utf-8\">"
            '<meta name="viewport" content="width=device-width,initial-scale=1">'
            "<title>WeChat OA draft preview</title></head>"
            '<body style="margin:0;background:#f2f2f2;">'
            '<main style="box-sizing:border-box;max-width:677px;margin:0 auto;'
            'padding:28px 20px 60px;background:#ffffff;">'
            '<div style="font-size:24px;line-height:1.4;font-weight:700;'
            'color:#222222;margin-bottom:24px;">'
            f"{html.escape(title)}</div>{content}</main></body></html>"
        )


def _placeholder(index: int) -> str:
    return f"wxcli-image-{index:03d}"


def _formatted_text(paragraph: Paragraph) -> str:
    if not paragraph.runs:
        return html.escape(paragraph.text).replace("\n", "<br />")
    values: list[str] = []
    for run in paragraph.runs:
        if not run.text:
            continue
        values.append(_formatted_run(run))
    return "".join(values) or html.escape(paragraph.text)


def _formatted_run(run: Run) -> str:
    value = html.escape(run.text).replace("\n", "<br />")
    styles: list[str] = []
    if run.bold:
        styles.append("font-weight:700")
    if run.italic:
        styles.append("font-style:italic")
    if run.underline:
        styles.append("text-decoration:underline")
    if run.font.color and run.font.color.rgb:
        styles.append(f"color:#{run.font.color.rgb}")
    return f'<span style="{";".join(styles)}">{value}</span>' if styles else value


def _prepare_jpeg(blob: bytes, target: Path, limit: int) -> PreparedImage:
    try:
        with Image.open(BytesIO(blob)) as opened:
            opened.load()
            transposed = ImageOps.exif_transpose(opened)
            if transposed.mode in {"RGBA", "LA"}:
                rgba = transposed.convert("RGBA")
                image = Image.new("RGB", rgba.size, "white")
                image.paste(rgba, mask=rgba.getchannel("A"))
            else:
                image = transposed.convert("RGB")
    except (OSError, ValueError, UnidentifiedImageError) as error:
        raise WxcliError(ErrorCode.PARSING_ERROR, "An article image could not be decoded.") from error

    if image.width <= 0 or image.height <= 0:
        raise ValidationError("An article image has invalid dimensions.")
    prepared = _jpeg_bytes_under_limit(image, limit)
    try:
        target.write_bytes(prepared)
    except OSError as error:
        raise WxcliError(
            ErrorCode.LOCAL_CONFIGURATION_ERROR,
            "The draft preview image could not be written.",
        ) from error
    return PreparedImage(
        path=target,
        original_bytes=len(blob),
        prepared_bytes=len(prepared),
        width=image.width,
        height=image.height,
        sha256=hashlib.sha256(prepared).hexdigest(),
    )


def _jpeg_bytes_under_limit(image: Image.Image, limit: int) -> bytes:
    working = image.copy()
    while working.width >= 320 and working.height >= 180:
        for quality in range(94, 49, -4):
            buffer = BytesIO()
            working.save(
                buffer,
                format="JPEG",
                quality=quality,
                optimize=True,
                progressive=True,
                subsampling="4:2:0",
            )
            value = buffer.getvalue()
            if len(value) < limit:
                return value
        next_size = (max(1, int(working.width * 0.88)), max(1, int(working.height * 0.88)))
        working = working.resize(next_size, Image.Resampling.LANCZOS)
    raise ValidationError("An image could not be compressed below the WeChat size limit.")


def _validate_file(path: Path, suffix: str, name: str) -> Path:
    try:
        resolved = path.expanduser().resolve(strict=True)
    except OSError as error:
        raise ValidationError(f"The {name} does not exist.") from error
    if not resolved.is_file() or resolved.suffix.casefold() != suffix:
        raise ValidationError(f"The {name} must be a {suffix} file.")
    return resolved


def _validate_image_file(path: Path, name: str) -> Path:
    try:
        resolved = path.expanduser().resolve(strict=True)
    except OSError as error:
        raise ValidationError(f"The {name} does not exist.") from error
    if not resolved.is_file() or resolved.suffix.casefold() not in {".jpg", ".jpeg", ".png"}:
        raise ValidationError(f"The {name} must be a JPG or PNG file.")
    return resolved


def _prepare_empty_directory(path: Path) -> None:
    try:
        if path.exists() and any(path.iterdir()):
            raise ValidationError("The preview output directory must be empty.")
        path.mkdir(parents=True, exist_ok=True)
    except ValidationError:
        raise
    except OSError as error:
        raise WxcliError(
            ErrorCode.LOCAL_CONFIGURATION_ERROR,
            "The preview output directory is not writable.",
        ) from error


def _validate_empty_output_directory(path: Path) -> None:
    try:
        if path.exists() and (not path.is_dir() or any(path.iterdir())):
            raise ValidationError("The preview output directory must be empty.")
    except ValidationError:
        raise
    except OSError as error:
        raise WxcliError(
            ErrorCode.LOCAL_CONFIGURATION_ERROR,
            "The preview output directory could not be inspected.",
        ) from error


def _clean_failed_output(path: Path, *, remove_root: bool) -> None:
    """Remove only files created after the output directory passed the empty check."""
    try:
        for filename in ("preview.html", "manifest.json", "package.json"):
            (path / filename).unlink(missing_ok=True)
        image_dir = path / "images"
        if image_dir.exists():
            shutil.rmtree(image_dir)
        if remove_root:
            path.rmdir()
    except OSError:
        # Preserve the original conversion error if local cleanup itself fails.
        return


def _optional_limited_text(value: str | None, limit: int, name: str) -> str | None:
    if value is None:
        return None
    normalized = value.strip()
    if not normalized:
        return None
    if len(normalized) > limit:
        raise ValidationError(f"The {name} must not exceed {limit} characters.")
    return normalized


def _validate_content(content: str) -> None:
    if len(content) >= _CONTENT_CHARACTER_LIMIT:
        raise ValidationError("The converted draft body must contain fewer than 20000 characters.")
    if len(content.encode("utf-8")) >= _CONTENT_BYTE_LIMIT:
        raise ValidationError("The converted draft body must be smaller than 1 MB.")


def _document_blocks(document: DocumentObject) -> list[Paragraph | Table]:
    values: list[Paragraph | Table] = []
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            values.append(Paragraph(child, document))
        elif child.tag == qn("w:tbl"):
            values.append(Table(child, document))
    return values


def _list_tag(paragraph: Paragraph, style_name: str) -> str | None:
    lowered = style_name.casefold()
    if "list bullet" in lowered:
        return "ul"
    if "list number" in lowered:
        return "ol"
    properties = paragraph._p.pPr
    if properties is None or properties.numPr is None:
        return None
    return "ol"


def _safe_link(value: object) -> str | None:
    if not isinstance(value, str):
        return None
    from urllib.parse import urlsplit

    try:
        parsed = urlsplit(value)
        port = parsed.port
    except ValueError:
        return None
    if (
        parsed.scheme.casefold() not in {"http", "https"}
        or not parsed.hostname
        or parsed.username is not None
        or parsed.password is not None
        or port is not None
    ):
        return None
    return value


def _package_value(draft: PreparedDraft, root: Path) -> dict[str, object]:
    def image_value(image: PreparedImage) -> dict[str, object]:
        return {
            "path": image.path.relative_to(root).as_posix(),
            "original_bytes": image.original_bytes,
            "prepared_bytes": image.prepared_bytes,
            "width": image.width,
            "height": image.height,
            "sha256": image.sha256,
        }

    return {
        "schema_version": 1,
        "package_sha256": draft.package_sha256,
        "title": draft.title,
        "author": draft.author,
        "digest": draft.digest,
        "content_template": draft.content_template,
        "images": [image_value(image) for image in draft.images],
        "cover": image_value(draft.cover),
        "preview": draft.preview.model_dump(mode="json"),
    }


def _load_prepared_image(root: Path, value: object) -> PreparedImage:
    if not isinstance(value, dict):
        raise TypeError
    relative = Path(str(value["path"]))
    if relative.is_absolute() or ".." in relative.parts:
        raise ValueError
    path = (root / relative).resolve(strict=True)
    try:
        path.relative_to(root)
    except ValueError as error:
        raise ValueError from error
    blob = path.read_bytes()
    digest = hashlib.sha256(blob).hexdigest()
    if digest != value.get("sha256") or len(blob) != value.get("prepared_bytes"):
        raise ValueError
    return PreparedImage(
        path=path,
        original_bytes=int(value["original_bytes"]),
        prepared_bytes=len(blob),
        width=int(value["width"]),
        height=int(value["height"]),
        sha256=digest,
    )


def _string_or_none(value: object) -> str | None:
    if value is None:
        return None
    if not isinstance(value, str):
        raise TypeError
    return value
